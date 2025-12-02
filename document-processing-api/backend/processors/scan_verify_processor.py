# backend/processors/scan_verify_processor.py


from pathlib import Path
from docx import Document
import pandas as pd
from .base_processor import BaseProcessor
import logging

logger = logging.getLogger(__name__)

class ScanVerifyProcessor(BaseProcessor):
    """Scan and verify Word documents"""
    
    async def process(self, file_path: Path, output_dir: Path) -> dict:
        """Extract document structure and verify content"""
        
        logger.info(f"Scanning and verifying {file_path.name}")
        
        try:
            # Use python-docx for structure analysis
            doc = Document(file_path)
            
            analysis = {
                "filename": file_path.name,
                "word_count": 0,
                "sections": [],
                "images": 0,
                "tables": len(doc.tables),
                "paragraphs": len(doc.paragraphs),
                "verification": {}
            }
            
            # Extract sections and content
            current_section = None
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                    
                # Check for section headers (using style)
                if para.style and para.style.name and 'Heading' in para.style.name:
                    current_section = {
                        "title": text,
                        "level": para.style.name,
                        "content": [],
                        "word_count": 0
                    }
                    analysis["sections"].append(current_section)
                elif current_section:
                    current_section["content"].append(text)
                    word_count = len(text.split())
                    current_section["word_count"] += word_count
                    analysis["word_count"] += word_count
                else:
                    # Content before first heading
                    analysis["word_count"] += len(text.split())
            
            # Count images through relationships
            try:
                for rel in doc.part.rels.values():
                    if "image" in rel.target_ref:
                        analysis["images"] += 1
            except:
                pass
            
            # Extract table information
            table_info = []
            for i, table in enumerate(doc.tables):
                table_info.append({
                    "Table": f"Table {i+1}",
                    "Rows": len(table.rows),
                    "Columns": len(table.columns) if table.rows else 0
                })
            
            # Verification checks
            analysis["verification"] = {
                "has_title": len(analysis["sections"]) > 0,
                "has_content": analysis["word_count"] > 100,
                "structure_valid": len(analysis["sections"]) > 1,
                "has_tables": analysis["tables"] > 0,
                "has_images": analysis["images"] > 0,
                "ready_for_conversion": True
            }
            
            # Save analysis to Excel
            output_filename = f"{file_path.stem}_analysis.xlsx"
            output_path = output_dir / output_filename
            
            with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
                # Summary sheet
                summary_data = pd.DataFrame({
                    "Metric": [
                        "Word Count", 
                        "Sections", 
                        "Paragraphs",
                        "Images", 
                        "Tables", 
                        "Ready for Conversion"
                    ],
                    "Value": [
                        analysis["word_count"],
                        len(analysis["sections"]),
                        analysis["paragraphs"],
                        analysis["images"],
                        analysis["tables"],
                        "Yes" if analysis["verification"]["ready_for_conversion"] else "No"
                    ]
                })
                summary_data.to_excel(writer, sheet_name="Summary", index=False)
                
                # Sections detail
                if analysis["sections"]:
                    sections_data = []
                    for section in analysis["sections"]:
                        sections_data.append({
                            "Section": section["title"],
                            "Level": section["level"],
                            "Word Count": section["word_count"],
                            "Content Preview": section["content"][0][:100] if section["content"] else ""
                        })
                    pd.DataFrame(sections_data).to_excel(writer, sheet_name="Sections", index=False)
                
                # Tables detail
                if table_info:
                    pd.DataFrame(table_info).to_excel(writer, sheet_name="Tables", index=False)
                
                # Verification details
                verification_data = pd.DataFrame({
                    "Check": list(analysis["verification"].keys()),
                    "Result": list(analysis["verification"].values())
                })
                verification_data.to_excel(writer, sheet_name="Verification", index=False)
            
            logger.info(f"Analysis saved to: {output_path}")
            
            return {
                "filename": file_path.name,
                "output_filename": output_filename,
                "path": str(output_path),
                "analysis": analysis,
                "success": True
            }
            
        except Exception as e:
            logger.error(f"Failed to analyze {file_path.name}: {str(e)}")
            raise

